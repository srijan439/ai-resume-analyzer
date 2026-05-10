"""
Resume parsing utilities.

This module owns file validation, text extraction, and simple profile parsing.
Keeping this code outside Streamlit makes it reusable from tests, scripts, or a
future API backend.
"""

import re
from pathlib import Path
from typing import Dict, List

import docx
import pdfplumber


ALLOWED_EXTENSIONS = {".pdf", ".docx"}


def validate_resume_file(file_name: str) -> bool:
    """Return True when the uploaded file has a supported extension."""
    return Path(file_name).suffix.lower() in ALLOWED_EXTENSIONS


def parse_resume(file_path: str) -> str:
    """
    Extract text from a PDF or DOCX resume.

    Args:
        file_path: Local path of the uploaded resume.

    Returns:
        Cleaned resume text. An empty string means extraction failed.
    """
    path = Path(file_path)
    extension = path.suffix.lower()

    if extension == ".pdf":
        raw_text = extract_text_from_pdf(path)
    elif extension == ".docx":
        raw_text = extract_text_from_docx(path)
    else:
        raise ValueError("Unsupported file type. Please upload a PDF or DOCX file.")

    return clean_text(raw_text)


def extract_text_from_pdf(file_path: Path) -> str:
    """Extract readable text from every page of a PDF using pdfplumber."""
    page_text: List[str] = []

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            page_text.append(text)

    return "\n".join(page_text)


def extract_text_from_docx(file_path: Path) -> str:
    """Extract paragraph and table text from a DOCX resume."""
    document = docx.Document(file_path)
    parts: List[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def clean_text(text: str) -> str:
    """Normalize whitespace while keeping line breaks useful for section parsing."""
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_contact_details(resume_text: str) -> Dict[str, str]:
    """Extract basic contact fields with readable regular expressions."""
    email_match = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", resume_text)
    phone_match = re.search(
        r"(?:(?:\+?\d{1,3}[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?)?\d{3}[\s.-]?\d{4})",
        resume_text,
    )

    return {
        "name": extract_name(resume_text),
        "email": email_match.group(0) if email_match else "Not found",
        "phone": phone_match.group(0) if phone_match else "Not found",
    }


def extract_name(resume_text: str) -> str:
    """
    Estimate the candidate name from the first meaningful line.

    This beginner-friendly heuristic avoids pretending we have perfect identity
    extraction. It works well for common resumes where the name is the first line.
    """
    ignored_words = {"resume", "curriculum vitae", "cv"}

    for line in resume_text.splitlines()[:8]:
        candidate = line.strip()
        if not candidate or candidate.lower() in ignored_words:
            continue
        if "@" in candidate or re.search(r"\d", candidate):
            continue
        if 2 <= len(candidate.split()) <= 4:
            return candidate

    return "Not found"


def extract_resume_sections(resume_text: str) -> Dict[str, str]:
    """Collect common resume sections into a dictionary for dashboard display."""
    section_names = [
        "summary",
        "profile",
        "skills",
        "technical skills",
        "experience",
        "work experience",
        "projects",
        "education",
        "certifications",
    ]
    sections = {name: "" for name in section_names}
    lines = resume_text.splitlines()
    current_section = ""

    for line in lines:
        normalized = line.strip().lower().rstrip(":")
        if normalized in sections:
            current_section = normalized
            continue
        if current_section and line.strip():
            sections[current_section] += line.strip() + "\n"

    return {key: value.strip() for key, value in sections.items() if value.strip()}


def extract_profile_summary(resume_text: str, skills: List[str]) -> Dict[str, object]:
    """
    Return the exact resume fields requested by the project scope.

    The UI can call this one function when it needs a beginner-friendly profile
    object containing name, email, phone, skills, education, and experience.
    """
    contact_details = extract_contact_details(resume_text)
    sections = extract_resume_sections(resume_text)

    education = sections.get("education", "Not found")
    experience = sections.get("experience") or sections.get("work experience") or "Not found"

    return {
        "name": contact_details["name"],
        "email": contact_details["email"],
        "phone": contact_details["phone"],
        "skills": skills,
        "education": education,
        "experience": experience,
    }
